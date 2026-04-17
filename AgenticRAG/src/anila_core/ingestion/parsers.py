"""Document parsers for the ANILA Core RAG pipeline.

Supports: .txt, .md, .pdf, .docx, .doc, .odt

Each parser implements the DocumentParser Protocol:
    parse(file_path: str) -> ParsedDocument

ParserRegistry selects the correct parser by file extension.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Protocol, runtime_checkable

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Result of parsing a document file."""

    content: str
    metadata: dict = field(default_factory=dict)
    source_path: str = ""
    format: str = ""


@runtime_checkable
class DocumentParser(Protocol):
    """Protocol for document parsers."""

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a document file and return its text content with metadata."""
        ...


class PlainTextParser:
    """Parser for .txt files."""

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        content = path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(
            content=content,
            metadata={"title": path.stem},
            source_path=file_path,
            format="txt",
        )


class MarkdownParser:
    """Parser for .md files. Preserves heading structure in metadata."""

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        content = path.read_text(encoding="utf-8", errors="replace")

        # Extract top-level heading as title
        title = path.stem
        for line in content.splitlines():
            if line.startswith("# "):
                title = line[2:].strip()
                break

        return ParsedDocument(
            content=content,
            metadata={"title": title},
            source_path=file_path,
            format="md",
        )


class PdfParser:
    """Parser for .pdf files using pymupdf4llm (PDF → Markdown)."""

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            import pymupdf4llm  # type: ignore[import]
        except ImportError as exc:
            raise ImportError(
                "pymupdf4llm is required for PDF parsing. "
                "Install with: pip install 'anila-core[rag]'"
            ) from exc

        path = Path(file_path)
        md_content: str = pymupdf4llm.to_markdown(file_path)
        return ParsedDocument(
            content=md_content,
            metadata={"title": path.stem, "pages": _count_pdf_pages(file_path)},
            source_path=file_path,
            format="pdf",
        )


def _count_pdf_pages(file_path: str) -> int:
    """Return page count for a PDF, or 0 on failure."""
    try:
        import fitz  # type: ignore[import]  # pymupdf

        doc = fitz.open(file_path)
        count = doc.page_count
        doc.close()
        return count
    except Exception:
        return 0


class DocxParser:
    """Parser for .docx files using python-docx."""

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            from docx import Document  # type: ignore[import]
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install 'anila-core[rag]'"
            ) from exc

        path = Path(file_path)
        doc = Document(file_path)

        parts: list[str] = []
        title = path.stem

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Preserve heading structure as Markdown
            if para.style.name.startswith("Heading 1"):
                parts.append(f"# {text}")
                if title == path.stem:
                    title = text
            elif para.style.name.startswith("Heading 2"):
                parts.append(f"## {text}")
            elif para.style.name.startswith("Heading 3"):
                parts.append(f"### {text}")
            else:
                parts.append(text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))

        return ParsedDocument(
            content="\n\n".join(parts),
            metadata={"title": title},
            source_path=file_path,
            format="docx",
        )


class DocParser:
    """Parser for legacy .doc files using antiword CLI."""

    def parse(self, file_path: str) -> ParsedDocument:
        import subprocess

        path = Path(file_path)
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True,
                text=True,
                timeout=30,
            )
            content = result.stdout
            if not content and result.returncode != 0:
                raise RuntimeError(f"antiword failed: {result.stderr}")
        except FileNotFoundError as exc:
            raise RuntimeError(
                "antiword is required for .doc parsing. Install with: apt install antiword"
            ) from exc

        return ParsedDocument(
            content=content,
            metadata={"title": path.stem},
            source_path=file_path,
            format="doc",
        )


class OdtParser:
    """Parser for .odt files using odfpy."""

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            from odf.opendocument import load as odf_load  # type: ignore[import]
            from odf.teletype import extractText  # type: ignore[import]
        except ImportError as exc:
            raise ImportError(
                "odfpy is required for ODT parsing. "
                "Install with: pip install 'anila-core[rag]'"
            ) from exc

        path = Path(file_path)
        doc = odf_load(file_path)

        parts: list[str] = []
        title = path.stem

        for element in doc.text.childNodes:
            text = extractText(element).strip()
            if not text:
                continue
            tag = element.qname[1] if hasattr(element, "qname") else ""
            if tag == "h":
                outline = element.getAttribute("text:outline-level") or "1"
                prefix = "#" * int(outline)
                parts.append(f"{prefix} {text}")
                if title == path.stem and int(outline) == 1:
                    title = text
            else:
                parts.append(text)

        return ParsedDocument(
            content="\n\n".join(parts),
            metadata={"title": title},
            source_path=file_path,
            format="odt",
        )


class ParserRegistry:
    """Select the appropriate parser based on file extension."""

    _PARSERS: dict[str, DocumentParser] = {
        ".txt": PlainTextParser(),
        ".md": MarkdownParser(),
        ".pdf": PdfParser(),
        ".docx": DocxParser(),
        ".doc": DocParser(),
        ".odt": OdtParser(),
    }

    @classmethod
    def get(cls, file_path: str) -> DocumentParser:
        """Return the parser for the given file extension."""
        ext = Path(file_path).suffix.lower()
        parser = cls._PARSERS.get(ext)
        if parser is None:
            supported = ", ".join(cls._PARSERS)
            raise ValueError(
                f"Unsupported file extension '{ext}'. Supported: {supported}"
            )
        return parser

    @classmethod
    def parse(cls, file_path: str) -> ParsedDocument:
        """Parse a document using the appropriate parser."""
        return cls.get(file_path).parse(file_path)

    @classmethod
    def supported_extensions(cls) -> list[str]:
        """Return list of supported file extensions."""
        return list(cls._PARSERS)
