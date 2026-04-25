"""Document parsers for the AgenticRAG pipeline.

Supports: .txt, .md, .pdf, .docx, .doc, .odt, .rtf, plus standalone
image files (.png, .jpg, .jpeg, .webp, .gif, .bmp).

Each parser implements the DocumentParser Protocol:

    parse(file_path: str) -> ParsedDocument

Parsers are synchronous and VLM-free. Images embedded inside a document
are extracted into ``ParsedDocument.images`` and a placeholder of the
form ``[[IMAGE:<id>]]`` is inserted at the corresponding position in
``content``. The IngestionService later asks a VisionProvider to caption
each image and rewrites the placeholders in-place.

ParserRegistry selects the correct parser by file extension.
"""

from __future__ import annotations

import logging
import mimetypes
import os
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Protocol, runtime_checkable

logger = logging.getLogger(__name__)


@dataclass
class ImageRef:
    """An image extracted from (or equal to) a document."""

    image_id: str
    image_bytes: bytes
    mime: str = "image/png"
    page: int | None = None       # 1-based page number for PDFs, None otherwise
    alt_text: str = ""             # alt / title attribute if available
    caption: str = ""               # filled in by the VLM step


@dataclass
class ParsedDocument:
    """Result of parsing a document file.

    ``content`` may contain placeholders of the form ``[[IMAGE:<id>]]``
    where ``<id>`` is a key in ``images``. The IngestionService resolves
    these placeholders before chunking.
    """

    content: str
    metadata: dict = field(default_factory=dict)
    source_path: str = ""
    format: str = ""
    images: dict[str, ImageRef] = field(default_factory=dict)


@runtime_checkable
class DocumentParser(Protocol):
    """Protocol for document parsers."""

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a document file and return its text content with metadata."""
        ...


def _new_image_id() -> str:
    return f"img_{uuid.uuid4().hex[:10]}"


# ──────────────────────────────────────────────────────────────────────
# Plain text / markdown / rtf
# ──────────────────────────────────────────────────────────────────────

class PlainTextParser:
    """Parser for .txt files."""

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        content = path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(
            content=content,
            metadata={"title": path.stem},
            source_path=file_path,
            format="txt",
        )


class MarkdownParser:
    """Parser for .md files. Preserves heading structure in metadata."""

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        content = path.read_text(encoding="utf-8", errors="replace")

        title = path.stem
        for line in content.splitlines():
            if line.startswith("# "):
                title = line[2:].strip()
                break

        return ParsedDocument(
            content=content,
            metadata={"title": title},
            source_path=file_path,
            format="md",
        )


class RtfParser:
    """Parser for .rtf files using striprtf (pure Python, no system deps)."""

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            from striprtf.striprtf import rtf_to_text  # type: ignore[import]
        except ImportError as exc:
            raise ImportError(
                "striprtf is required for RTF parsing. "
                "Install with: pip install 'agentic-rag[rag]'"
            ) from exc

        path = Path(file_path)
        raw = path.read_text(encoding="utf-8", errors="replace")
        content = rtf_to_text(raw, errors="ignore")

        return ParsedDocument(
            content=content,
            metadata={"title": path.stem},
            source_path=file_path,
            format="rtf",
        )


# ──────────────────────────────────────────────────────────────────────
# PDF (with embedded image extraction)
# ──────────────────────────────────────────────────────────────────────

class PdfParser:
    """Parser for .pdf files.

    Text is extracted with pymupdf4llm (PDF → Markdown). Embedded images
    are extracted with pymupdf (fitz) and appended at the end of each
    page's markdown as ``[[IMAGE:<id>]]`` placeholders so the
    HierarchicalChunker can attach them under the surrounding heading.

    When digital extraction yields little or mostly ``<?>`` placeholders
    (font-subsetted PDFs), an optional OCR backend is invoked. The
    backend is constructed lazily on first need from environment
    variables — see ``ingestion.ocr.build_ocr_backend_from_env``.
    """

    _ocr_backend = None  # type: ignore[var-annotated]
    _ocr_initialised = False

    @classmethod
    def _get_ocr_backend(cls):
        if not cls._ocr_initialised:
            from .ocr import build_ocr_backend_from_env
            cls._ocr_backend = build_ocr_backend_from_env()
            cls._ocr_initialised = True
        return cls._ocr_backend

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            import pymupdf4llm  # type: ignore[import]
            import fitz  # type: ignore[import]  # pymupdf
        except ImportError as exc:
            raise ImportError(
                "pymupdf4llm and pymupdf are required for PDF parsing. "
                "Install with: pip install 'agentic-rag[rag]'"
            ) from exc

        path = Path(file_path)

        # Per-page markdown lets us interleave image placeholders with text.
        page_chunks = pymupdf4llm.to_markdown(file_path, page_chunks=True)

        images: dict[str, ImageRef] = {}
        doc = fitz.open(file_path)
        try:
            parts: list[str] = []
            for pno, page_entry in enumerate(page_chunks, start=1):
                page_md = (
                    page_entry.get("text", "")
                    if isinstance(page_entry, dict)
                    else str(page_entry)
                )
                parts.append(page_md.rstrip())

                page = doc[pno - 1]
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    try:
                        image_bytes, mime = _extract_pdf_image(doc, xref)
                    except Exception as exc:
                        logger.warning(
                            "Skip image xref=%s on page %d: %s", xref, pno, exc
                        )
                        continue

                    img_id = _new_image_id()
                    images[img_id] = ImageRef(
                        image_id=img_id,
                        image_bytes=image_bytes,
                        mime=mime,
                        page=pno,
                    )
                    parts.append(f"\n[[IMAGE:{img_id}]]\n")
        finally:
            doc.close()

        # Page join uses ``\f\n`` (form-feed) so consumers that need
        # per-page boundaries (e.g. the central ingestion-worker's
        # ``pdf-page`` chunker) can split on the marker without
        # re-parsing the PDF. ``\f`` is rarely emitted by PDF text
        # extractors so the marker is safe to insert. Markdown-style
        # consumers ignore it as whitespace.
        content = "\f\n".join(p for p in parts if p)
        ocr_used = False

        # Optional OCR fallback for scanned / font-subsetted PDFs.
        backend = self._get_ocr_backend()
        if backend is not None:
            from .ocr import needs_ocr_fallback
            if needs_ocr_fallback(content):
                logger.info(
                    "PDF %s text extraction looks unusable — running OCR fallback",
                    path.name,
                )
                try:
                    ocr_text = backend.extract(file_path)
                    if ocr_text.strip():
                        content = ocr_text
                        ocr_used = True
                except Exception as exc:
                    logger.warning(
                        "OCR fallback failed for %s: %s — keeping native extraction",
                        path.name, exc,
                    )

        return ParsedDocument(
            content=content,
            metadata={
                "title": path.stem,
                "pages": len(page_chunks),
                "embedded_images": len(images),
                "ocr_used": ocr_used,
            },
            source_path=file_path,
            format="pdf",
            images=images,
        )


def _extract_pdf_image(doc, xref: int) -> tuple[bytes, str]:
    """Return (bytes, mime) for a PDF image xref, preferring the native format."""
    import fitz  # type: ignore[import]

    info = doc.extract_image(xref)
    if info and info.get("image"):
        ext = (info.get("ext") or "png").lower()
        mime = f"image/{'jpeg' if ext == 'jpg' else ext}"
        return info["image"], mime

    pix = fitz.Pixmap(doc, xref)
    try:
        if pix.n - pix.alpha >= 4:  # CMYK / DeviceN → convert to RGB
            pix = fitz.Pixmap(fitz.csRGB, pix)
        return pix.tobytes("png"), "image/png"
    finally:
        pix = None


# ──────────────────────────────────────────────────────────────────────
# DOCX (with embedded image extraction)
# ──────────────────────────────────────────────────────────────────────

class DocxParser:
    """Parser for .docx files using python-docx.

    Paragraphs are emitted as Markdown with heading levels preserved.
    Embedded images are extracted per paragraph (when a run contains a
    ``w:drawing`` element) and inserted inline as ``[[IMAGE:<id>]]``.
    Tables are rendered as pipe-separated rows.
    """

    _NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    _NS_R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            from docx import Document  # type: ignore[import]
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install 'agentic-rag[rag]'"
            ) from exc

        path = Path(file_path)
        doc = Document(file_path)

        parts: list[str] = []
        images: dict[str, ImageRef] = {}
        title = path.stem

        for para in doc.paragraphs:
            text = para.text.strip()
            para_images = self._collect_para_images(para, doc, images)

            if not text and not para_images:
                continue

            style = para.style.name if para.style is not None else ""
            if style.startswith("Heading 1"):
                parts.append(f"# {text}")
                if title == path.stem and text:
                    title = text
            elif style.startswith("Heading 2"):
                parts.append(f"## {text}")
            elif style.startswith("Heading 3"):
                parts.append(f"### {text}")
            elif style.startswith("Heading 4"):
                parts.append(f"#### {text}")
            elif text:
                parts.append(text)

            for img_id in para_images:
                parts.append(f"[[IMAGE:{img_id}]]")

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))

        return ParsedDocument(
            content="\n\n".join(parts),
            metadata={"title": title, "embedded_images": len(images)},
            source_path=file_path,
            format="docx",
            images=images,
        )

    def _collect_para_images(
        self,
        paragraph,
        doc,
        images: dict[str, ImageRef],
    ) -> list[str]:
        """Find inline images inside a paragraph and register them."""
        ids: list[str] = []
        for run in paragraph.runs:
            for blip in run._element.iter(f"{self._NS_A}blip"):
                rid = blip.get(f"{self._NS_R}embed")
                if not rid:
                    continue
                try:
                    image_part = doc.part.related_parts[rid]
                    image_bytes = image_part.blob
                    mime = getattr(image_part, "content_type", "image/png")
                except Exception as exc:
                    logger.warning("DOCX image rId=%s extraction failed: %s", rid, exc)
                    continue

                img_id = _new_image_id()
                images[img_id] = ImageRef(
                    image_id=img_id,
                    image_bytes=image_bytes,
                    mime=mime,
                )
                ids.append(img_id)
        return ids


# ──────────────────────────────────────────────────────────────────────
# Legacy .doc (antiword)
# ──────────────────────────────────────────────────────────────────────

class DocParser:
    """Parser for legacy .doc files using antiword CLI."""

    def parse(self, file_path: str) -> ParsedDocument:
        import subprocess

        path = Path(file_path)
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True,
                text=True,
                timeout=30,
            )
            content = result.stdout
            if not content and result.returncode != 0:
                raise RuntimeError(f"antiword failed: {result.stderr}")
        except FileNotFoundError as exc:
            raise RuntimeError(
                "antiword is required for .doc parsing. Install with: apt install antiword"
            ) from exc

        return ParsedDocument(
            content=content,
            metadata={"title": path.stem},
            source_path=file_path,
            format="doc",
        )


# ──────────────────────────────────────────────────────────────────────
# ODT
# ──────────────────────────────────────────────────────────────────────

class OdtParser:
    """Parser for .odt files using odfpy."""

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            from odf.opendocument import load as odf_load  # type: ignore[import]
            from odf.teletype import extractText  # type: ignore[import]
        except ImportError as exc:
            raise ImportError(
                "odfpy is required for ODT parsing. "
                "Install with: pip install 'agentic-rag[rag]'"
            ) from exc

        path = Path(file_path)
        doc = odf_load(file_path)

        parts: list[str] = []
        title = path.stem

        for element in doc.text.childNodes:
            text = extractText(element).strip()
            if not text:
                continue
            tag = element.qname[1] if hasattr(element, "qname") else ""
            if tag == "h":
                outline = element.getAttribute("text:outline-level") or "1"
                prefix = "#" * int(outline)
                parts.append(f"{prefix} {text}")
                if title == path.stem and int(outline) == 1:
                    title = text
            else:
                parts.append(text)

        return ParsedDocument(
            content="\n\n".join(parts),
            metadata={"title": title},
            source_path=file_path,
            format="odt",
        )


# ──────────────────────────────────────────────────────────────────────
# Standalone image files
# ──────────────────────────────────────────────────────────────────────

_IMAGE_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
    ".bmp": "image/bmp",
}


class ImageParser:
    """Parser for standalone image files.

    The parsed document contains a single ``[[IMAGE:<id>]]`` placeholder
    which the IngestionService turns into a VLM caption.
    """

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        ext = path.suffix.lower()
        mime = (
            _IMAGE_MIME.get(ext)
            or mimetypes.guess_type(file_path)[0]
            or "application/octet-stream"
        )
        image_bytes = path.read_bytes()

        img_id = _new_image_id()
        images = {
            img_id: ImageRef(
                image_id=img_id,
                image_bytes=image_bytes,
                mime=mime,
            )
        }

        return ParsedDocument(
            content=f"# {path.stem}\n\n[[IMAGE:{img_id}]]",
            metadata={"title": path.stem, "embedded_images": 1},
            source_path=file_path,
            format="image",
            images=images,
        )


# ──────────────────────────────────────────────────────────────────────
# Registry
# ──────────────────────────────────────────────────────────────────────

class ParserRegistry:
    """Select the appropriate parser based on file extension.

    By default uses the lightweight native parsers (pymupdf4llm, python-docx,
    odfpy, ...). Set ``DOC_PARSER=docling`` in the environment to route
    PDF/DOCX/PPTX/XLSX/HTML through Docling instead — see
    ``ingestion.docling_parser`` for the trade-offs.
    """

    _PARSERS: dict[str, DocumentParser] = {
        ".txt": PlainTextParser(),
        ".md": MarkdownParser(),
        ".rtf": RtfParser(),
        ".pdf": PdfParser(),
        ".docx": DocxParser(),
        ".doc": DocParser(),
        ".odt": OdtParser(),
        # standalone images
        ".png": ImageParser(),
        ".jpg": ImageParser(),
        ".jpeg": ImageParser(),
        ".webp": ImageParser(),
        ".gif": ImageParser(),
        ".bmp": ImageParser(),
    }

    _docling_parser = None  # type: ignore[var-annotated]
    _docling_initialised = False

    @classmethod
    def _get_docling_parser(cls):
        """Return the cached Docling parser, or None when DOC_PARSER!=docling.

        The env var is consulted on every call (cheap) so switching it off
        immediately takes effect, but the heavy DoclingParser instance is
        constructed at most once per process.
        """
        if os.getenv("DOC_PARSER", "native").lower() != "docling":
            return None
        if not cls._docling_initialised:
            try:
                from .docling_parser import build_docling_parser_from_env
                cls._docling_parser = build_docling_parser_from_env()
            except Exception as exc:  # pragma: no cover - defensive
                logger.warning(
                    "Docling parser construction failed: %s — falling back to native",
                    exc,
                )
                cls._docling_parser = None
            cls._docling_initialised = True
        return cls._docling_parser

    @classmethod
    def _reset_docling_cache(cls) -> None:
        """Test hook: drop the cached Docling parser so env changes re-apply."""
        cls._docling_parser = None
        cls._docling_initialised = False

    @classmethod
    def get(cls, file_path: str) -> DocumentParser:
        """Return the parser for the given file extension."""
        ext = Path(file_path).suffix.lower()

        # Docling override for the formats it actually handles.
        from .docling_parser import DOCLING_SUPPORTED_EXTS  # local import: cheap, no model load
        if ext in DOCLING_SUPPORTED_EXTS:
            docling = cls._get_docling_parser()
            if docling is not None:
                return docling

        parser = cls._PARSERS.get(ext)
        if parser is None:
            supported = ", ".join(sorted(cls._PARSERS))
            raise ValueError(
                f"Unsupported file extension '{ext}'. Supported: {supported}"
            )
        return parser

    @classmethod
    def parse(cls, file_path: str) -> ParsedDocument:
        """Parse a document using the appropriate parser."""
        return cls.get(file_path).parse(file_path)

    @classmethod
    def supported_extensions(cls) -> list[str]:
        """Return list of supported file extensions."""
        return list(cls._PARSERS)
